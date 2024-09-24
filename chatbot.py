import streamlit as st
import PyPDF2
import docx
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import FAISS
from langchain.chains import ConversationalRetrievalChain
from langchain_openai import ChatOpenAI
from langchain.schema import Document
from string import Template
import re
import os
import time
from langchain.prompts import ChatPromptTemplate
from funcs import load_css, load_local_font
# 페이지 로드
from info import info_page 
from form import form_page
import pandas as pd 

def chat_page():
    # PDF 파일 읽기 함수
    def read_pdf(file_path):
        pdf_text = []
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            for page_num in range(len(reader.pages)):
                page = reader.pages[page_num]
                pdf_text.append(page.extract_text())
        return pdf_text

    # DOCX 파일 읽기 함수
    def read_docx(file_path):
        doc = docx.Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return '\n'.join(full_text)

    # 파일 경로
    pdf_file_path = "./data/docx/02-우리_몸이_원하는_삼삼한밥상_Ⅸ_본문(화면용_펼침).pdf"
    docx_file_path = "./data/docx/recipes.docx"
    pdf_file_path2 = './data/docx/1.pdf'

    # PDF 및 DOCX 파일에서 텍스트 불러오기
    pdf_content = read_pdf(pdf_file_path)
    docx_content = read_docx(docx_file_path)
    pdf_content2 = read_pdf(pdf_file_path2)



    # 데이터를 딕셔너리 형태로 정리
    data = {
        "항목": [
            "이름", "만나이", "성별", "키", "체중", "허리둘레", 
            "이상지질혈증 여부", "당뇨병 여부", 
            "음주 여부", "음주 빈도", "1회 주량", 
            "폭음 횟수", "절주 권고", 
            "음주 상담 유무",
            "고강도 운동 여부", "1주일", "1회 운동 시간", 
            "중강도 운동 여부", "1주일", "1회 운동 시간?", 
            "걷기, 자전거 운동", "1주일", "총 운동 시간",
            "음주 점수", "신체활동 점수", "고혈압 확률",
        ],
        "값": [
            "김첨지", 38, "남자", "173.0 cm", "89.0 kg", "88.9 cm",
            "없음", "없음",
            "마신다", "월1회정도", "1-2잔",
            "월 1회 미만", "없음", "없음",
            "안 한다", "None", "None", 
            "안 한다", "None", "None",
            "한다", "5일", "2시간 0분",
            "9점", "6.39859점", "21.41%",
            
        ]
    }

    # 데이터프레임 생성

    df = pd.DataFrame(data)
    df1 = df.reset_index(drop=True)

    # 사이드바에 데이터프레임 표시
    st.sidebar.title("김첨지님의 건강 보고서")
    st.sidebar.dataframe(df1)

    # PDF 및 DOCX 데이터를 문서 리스트로 변환
    documents = []

    # pdf_content는 리스트 형태일 가능성이 높으므로 개별 텍스트로 변환
    for text in pdf_content:
        if isinstance(text, str):
            documents.append(Document(page_content=text))
        else:
            st.error("PDF 내용이 올바른 문자열 형식이 아닙니다.")

    # docx_content는 하나의 문자열이므로 바로 추가
    if isinstance(docx_content, str):
        documents.append(Document(page_content=docx_content))
    else:
        st.error("DOCX 내용이 올바른 문자열 형식이 아닙니다.")

    # pdf_content2도 리스트일 가능성이 있으므로 개별 텍스트로 변환
    for text in pdf_content2:
        if isinstance(text, str):
            documents.append(Document(page_content=text))
        else:
            st.error("두 번째 PDF 내용이 올바른 문자열 형식이 아닙니다.")

    # 텍스트를 청크로 분리하는 함수
    def get_text_chunks(documents):
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=512,  # 청크 크기를 더 줄임
            chunk_overlap=50,
            length_function=len
        )
        chunks = text_splitter.split_documents(documents)
        return chunks

    # 벡터 스토어 생성 함수
    def get_vectorstore(_text_chunks):
        try:
            # Try using 'cuda' device first
            embeddings = HuggingFaceEmbeddings(
                model_name="jhgan/ko-sroberta-multitask",
                model_kwargs={'device': 'mps'},
                encode_kwargs={'normalize_embeddings': True}
            )
        except RuntimeError:
            # If there's an error, fall back to 'mps' device
            embeddings = HuggingFaceEmbeddings(
                model_name="jhgan/ko-sroberta-multitask",
                model_kwargs={'device': 'mps'},
                encode_kwargs={'normalize_embeddings': True}
            )
        
        vectordb = FAISS.from_documents(_text_chunks, embeddings)
        return vectordb

    # 시스템 프롬프트 템플릿 생성
    system_message = """
    data = {
        "항목": [
            "이름", "만나이", "성별", "키", "체중", "허리둘레", 
            "이상지질혈증 여부", "당뇨병 여부", 
            "술을 마십니까?", "술을 얼마나 자주 마십니까?", "한 번에 술을 얼마나 마십니까?", 
            "한 번의 술자리에서 7잔 이상을 마시는 횟수", "술을 끊거나 줄이라는 권고를 받은 적이 있습니까?", 
            "최근 1년 동안 음주 문제로 상담을 받아본 적이 있습니까?",
            "고강도 운동 여부", "1주일에 며칠 하십니까?", "한 번 할 때 몇 시간 하십니까?", 
            "중강도 운동 여부", "1주일에 며칠 하십니까?", "한 번 할 때 몇 시간 하십니까?", 
            "걷기나 자전거를 이용하십니까?", "1주일에 며칠 하십니까?", "대략 몇 시간 움직이십니까?",
            "음주 점수", "신체활동 점수", "고혈압 확률",
            "음주 점수 (동나이대/성별 평균)", "신체 활동 점수 (동나이대/성별 평균)", "고혈압 확률 (동나이대/성별 평균)"
        ],
        "값": [
            "김첨지", 38, "남자", "173.0 cm", "89.0 kg", "88.9 cm",
            "없음", "없음",
            "마신다", "월1회정도", "1-2잔",
            "월 1회 미만", "없음", "없음",
            "안 한다", "None", "None", 
            "안 한다", "None", "None",
            "한다", "5일", "2시간 0분",
            "9점", "6.39859점", "21.41%",
            "25.91점", "5.56점", "25.32%"
        ]
    }
    너는 이 데이터의 내용을 토대로 건강을 책임지는 영양 전문가야.

    """

    # 텍스트 청크 생성
    text_chunks = get_text_chunks(documents)

    # 벡터 스토어 생성
    vectorstore = get_vectorstore(text_chunks)

    # 대화형 체인 생성 함수
    def get_conversation_chain(_vectorstore):
        llm = ChatOpenAI(
            base_url="http://localhost:1234/v1",
            api_key="lm-studio",
            model="teddylee777/EEVE-Korean-Instruct-10.8B-v1.0-gguf",
            temperature=0.3,
            streaming=True
        )

        retriever = _vectorstore.as_retriever(
            search_type='similarity',
            search_kwargs={"k": 10},
            verbose=False
        )

        # 시스템 메시지를 포함하는 프롬프트 템플릿 생성
        prompt_template = ChatPromptTemplate.from_messages([
            ("system", system_message),
            ("user", "{question}")
        ])

        conversation_chain = ConversationalRetrievalChain.from_llm(
            llm=llm, 
            chain_type="stuff",
            retriever=retriever, 
            return_source_documents=True,

            verbose=False
        )
        return conversation_chain

    # 대화형 체인 생성
    conversation_chain = get_conversation_chain(vectorstore)

    if 'conversation' not in st.session_state:
        st.session_state.conversation = conversation_chain

    def get_prompt_template(query):
        # "식단"이 포함된 경우
        if "식단" in query:
            prompt = Template(f"""
            
            사용자가 원하는 기간이 있으면 그 기간, 없으면 하루
            {pdf_content2}의 내용을 토대로 대답해 줘. 
            {pdf_content2}에서 고혈압 확률이 20%가 넘으면 나트륨 1000mg 이하의 음식을
            무조건  {docx_file_path} 자료에서만 찾아줘.
            아니면 나트륨 제한없이 모든 음식을 
            무조건 {docx_file_path} 자료에서만 찾아줘.
            """)
        # "레시피"가 포함된 경우
        elif "레시피" in query:
            prompt = Template(f"""
            무조건 {docx_file_path} 자료에서 {query}의 레시피를 찾아서 자세히 설명해 줘.
            """)
        else:
            prompt = Template(f"""
            {pdf_file_path} 또는 {docx_file_path} 자료 또는 추가적인 정보를 포함해서 자유롭게 답변해 줘.
            {pdf_content2}의 내용을 토대로 대답해 줘.
            """)
        return prompt.substitute(query=query)

    # Chat logic
    if query := st.chat_input("식단 또는 레시피를 물어보세요."):
        # 입력된 query를 템플릿화
        formatted_query = get_prompt_template(query)

        with st.chat_message("user"):
            st.markdown(query)

        with st.chat_message("assistant"):
            chain = st.session_state.conversation

            with st.spinner("Thinking..."):
                # 빈 chat_history를 전달
                result = chain({"question": formatted_query, "chat_history": []})
                response = result['answer']
                source_documents = result['source_documents']

                response_container = st.empty()

                chunk_size = 5
                for i in range(0, len(response), chunk_size):
                    response_container.markdown(response[:i + chunk_size], unsafe_allow_html=True)
                    time.sleep(0.05)

                with st.expander("참고 문서 확인"):
                    for doc in source_documents:
                        st.markdown(doc.metadata.get('source', ''), help=doc.page_content)

    # 캐시 지우기 버튼
    if st.button('캐시 지우기'):
        st.cache_data.clear()
        st.cache_resource.clear()
        st.success("캐시가 지워졌습니다!")
